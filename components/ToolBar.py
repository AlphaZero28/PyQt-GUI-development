from PIL import Image
from PyQt5.QtWidgets import QAction, QFileDialog, QWidget, QToolBar,QSpinBox, QLabel, QPushButton
from PyQt5 import QtCore,Qt
from PyQt5.QtCore import QThread
from PyQt5.QtGui import QIcon, QImage
import fitz
# from matplotlib import style
from components.ImageProcessing import imgProcess
from components.config import DEBUG, WORK_ON_THREAD
from docx import Document
from docx.shared import Pt 
from components.Worker import SaveFileWorker, ocr_load_pages


class FileLoader(QtCore.QObject):
    finished = QtCore.pyqtSignal(list)
    progress = QtCore.pyqtSignal(list)
    def setPath(self,path):
        self.path = path
        print('path',path)

    def run(self):
        print('loading img')
        imgs = imgProcess.get_pages(self.path)
        # self.progress.emit(imgs)
        self.finished.emit(imgs)
                

class cToolBar(QWidget):
    def __init__(self, mainwindow):
        super(cToolBar, self).__init__()
        self.mainwindow = mainwindow
        self.zoom = 1
        self.total_page_number = 0

        # file toolbar
        fileToolBar = QToolBar("File", self)
        fileToolBar.setIconSize(QtCore.QSize(30, 30))

        openActionForToolBar = QAction(
            QIcon('./assets/file.png'), "&Open", self)
        openActionForToolBar.triggered.connect(self.openFiles)

        saveActionForToolBar = QAction(
            QIcon('./assets/save-solid.svg'), "&Save", self)
        saveActionForToolBar.triggered.connect(self.saveFiles)
        

        fileToolBar.addAction(openActionForToolBar)
        fileToolBar.addAction(saveActionForToolBar)

        # view toolbar
        viewToolBar = QToolBar("View", self)

        minusActionForToolBar = QAction(
            QIcon('./assets/minus.png'), "&Zoom Out", self)
        minusActionForToolBar.triggered.connect(self.zoom_out)

        plusActionForToolBar = QAction(
            QIcon('./assets/plus.png'), "&Zoom In", self)
        plusActionForToolBar.triggered.connect(self.zoom_in)

        viewToolBar.addAction(minusActionForToolBar)
        viewToolBar.addAction(plusActionForToolBar)



        # navigation toolbar
        # add label for navigation
        # label = QLabel()
        # label.setText('GoTo')
        navButton = QPushButton('GoTo')
        navButton.clicked.connect(self.navigationFunction)


        # add spinbox for navigation
        self.navigationBox = QSpinBox()
        # navigationBox.setFocusPolicy(Qt.NoFocus)
        # self.navigationBox.valueChanged.connect(self.navigationFunction)

        navigationBar = QToolBar("Navigation", self)
        navigationBar.addWidget(navButton)
        navigationBar.addWidget(self.navigationBox)        

        # add toolbar to mainwindow
        self.mainwindow.addToolBar(fileToolBar)
        self.mainwindow.addToolBar(viewToolBar)
        self.mainwindow.addToolBar(navigationBar)
        

    def navigationFunction(self):
        page_number = self.navigationBox.value()
        print(page_number)
        self.cmain_view.goto_specific_page(page_number)

    def set_main_view(self, cmain_view):
        self.cmain_view = cmain_view
        # self.get_pages('F:/My_Folder/__Projects__/InnovationGarage/Accessible-PDF-Reader/App/pyqt-pdfreader/Jhora Palok By Jibanananda Das (BDeBooks.Com)-pages-deleted.pdf')

    def zoom_in(self):
        self.zoom = self.zoom + 0.25
        self.cmain_view.set_zoom(self.zoom)

    def zoom_out(self):
        self.zoom = self.zoom - 0.25
        self.cmain_view.set_zoom(self.zoom)

    def openFiles(self):
        # if DEBUG:
        #     name = r'math-book-9-10-9-12.pdf'
        #     imgs = imgProcess.get_pages(name)
        #     self.cmain_view.set_imgs(imgs)
        #     return
        if DEBUG:
            fname = ['math-book-9-10-9-12.pdf']
        else:
            fname = QFileDialog.getOpenFileName(
                self, 'Open File', "", "PDF files (*.pdf);;")
        
        if fname[0]=='':
            return
        self.current_filename = fname[0]
        self.mainwindow.cstatus_bar.show_msg('Loading Page')
        # self.total_page_number = imgProcess.get_page_count(fname[0])
        self.cmain_view.set_path(fname[0])
        # # Step 2: Create a QThread object
        # self.thread = QtCore.QThread()
        # # Step 3: Create a worker object
        # self.worker = FileLoader()
        # self.worker.setPath(fname[0])
        # # Step 4: Move worker to the thread
        # self.worker.moveToThread(self.thread)
        # # Step 5: Connect signals and slots
        # self.thread.started.connect(self.worker.run)
        # self.worker.finished.connect(self.thread.quit)
        # self.worker.finished.connect(lambda imgs:self.cmain_view.set_imgs(imgs))
        # self.thread.finished.connect(self.thread.deleteLater)
        # # self.worker.progress.connect(add_img_to_view)
        # # Step 6: Start the thread
        # self.thread.start()

        # # Final resets
        # # self.longRunningBtn.setEnabled(False)
        # self.thread.finished.connect(
        #     lambda: print('finished')
        # )

        # # self.get_pages(fname[0])
        # self.imgs = imgProcess.get_pages(fname[0])
        # self.cmain_view.set_imgs(self.imgs)
        # # self.mainwindow.change_label(fname[0])


    def saveFiles(self):
        if self.total_page_number==0:
            return

        no_of_pages = self.total_page_number

        # file = str(QFileDialog.get)
        file = QFileDialog.getSaveFileName(self)
        filename = file[0]
        
        def save_as_docx(pages):
            # [pages] = lst
            # print(pages)
            # pages = []
            document = Document()
            style = document.styles['Normal']
            font = style.font
            font.name  = 'Arial'
            font.size = Pt(12)

            for page in pages:
                paragraph = document.add_paragraph(' ')
                paragraph.style = document.styles['Normal']
                # paragraph.add_run(page)
                # paragraph.add_run('\n')
                for line in page:
                    paragraph.add_run(line)
                    paragraph.add_run('\n')
                
                document.add_page_break()
            document.save(filename)
            self.mainwindow.cstatus_bar.show_msg('File Saved!')

        
        if WORK_ON_THREAD:
            # Step 2: Create a QThread object
            self.thread = QThread()
            # Step 3: Create a worker object
            self.worker = SaveFileWorker()
            self.worker.thread_function_init(self.current_filename,self.total_page_number,filename)
            # Step 4: Move worker to the thread
            self.worker.moveToThread(self.thread)
            # Step 5: Connect signals and slots
            self.thread.started.connect(self.worker.thread_function)
            self.worker.finished.connect(self.thread.quit)
            self.worker.finished.connect(self.worker.deleteLater)
            self.thread.finished.connect(self.thread.deleteLater)
            self.worker.progress.connect(save_as_docx)
            # Step 6: Start the thread
            self.thread.start()
        else:
            pages = ocr_load_pages(self.current_filename, self.total_page_number)
            save_as_docx(pages)

        # print(type(file))

        


        # with open('sample.txt', 'a') as f:
        #     for line in txt:
        #         f.write('\n')
        #         f.write(line)




